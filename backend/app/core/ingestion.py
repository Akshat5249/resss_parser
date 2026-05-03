import io
import re
import logging
import pdfplumber
import fitz  # PyMuPDF
from docx import Document
from fastapi import HTTPException
from typing import List, Optional

from app.shared.constants import SUPPORTED_FILE_TYPES, MAX_FILE_SIZE_MB, CHUNK_SIZE, CHUNK_OVERLAP

logger = logging.getLogger(__name__)

def validate_file(filename: str, content: bytes) -> None:
    """
    Checks if the file extension is supported and if the file size is within limits.
    """
    ext = "." + filename.split(".")[-1].lower()
    if ext not in SUPPORTED_FILE_TYPES:
        raise HTTPException(status_code=400, detail=f"File type {ext} not supported. Use {SUPPORTED_FILE_TYPES}")
    
    file_size_mb = len(content) / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE_MB:
        raise HTTPException(status_code=400, detail=f"File size exceeds limit of {MAX_FILE_SIZE_MB}MB")

def parse_pdf(content: bytes) -> str:
    """
    Parses a PDF file from bytes. Uses pdfplumber primary and fitz as fallback.
    """
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}. Falling back to fitz.")

    # Fallback if text is empty or very short
    if len(text.strip()) < 100:
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
        except Exception as e:
            logger.error(f"fitz fallback also failed: {e}")
            raise HTTPException(status_code=500, detail="Failed to extract text from PDF")

    return text

def parse_docx(content: bytes) -> str:
    """
    Parses a DOCX file from bytes, including tables.
    """
    try:
        doc = Document(io.BytesIO(content))
        full_text = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)
                        
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to extract text from DOCX")

def parse_file(content: bytes, filename: str) -> str:
    """
    Routes to the correct parser based on file extension.
    """
    ext = "." + filename.split(".")[-1].lower()
    if ext == ".pdf":
        return parse_pdf(content)
    elif ext == ".docx":
        return parse_docx(content)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file extension")

def clean_text(text: str) -> str:
    """
    Cleans extracted text by collapsing excessive whitespace and structure-preserving newlines.
    """
    # Remove 3+ consecutive newlines -> max 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove excessive spaces and tabs
    text = re.sub(r'[ \t]+', ' ', text)
    # Strip leading/trailing
    return text.strip()

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    Splits text into overlapping chunks, trying to preserve sentence boundaries.
    """
    # Simple sentence-based splitting (naive)
    sentences = re.split(r'(?<=[.!?]) +', text)
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= chunk_size:
            current_chunk += sentence + " "
        else:
            chunks.append(current_chunk.strip())
            # Overlap: take the last bit of the current chunk
            overlap_text = current_chunk[-(overlap):] if len(current_chunk) > overlap else current_chunk
            current_chunk = overlap_text + " " + sentence + " "
            
    if current_chunk:
        chunks.append(current_chunk.strip())
        
    return chunks

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        path = sys.argv[1]
        with open(path, "rb") as f:
            content = f.read()
            raw = parse_file(content, path)
            cleaned = clean_text(raw)
            print(f"--- First 500 chars of {path} ---\n")
            print(cleaned[:500])
    else:
        print("Usage: python ingestion.py <path_to_resume>")
